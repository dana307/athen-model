"""
DOCX renderer (python-docx).

Renders a ResearchReport as an institutional-style equity research note:
title block, recommendation banner, thesis, business/financial summary,
forecast, DCF, comparables, Monte Carlo, scenarios, risks, recommendation.
Charts are embedded if present in report.charts.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.schema import FIELD_LABELS

CR = 1e7
_REC_COLOR = {"BUY": RGBColor(0x1A, 0x7F, 0x37),
              "HOLD": RGBColor(0xB5, 0x8B, 0x00),
              "REDUCE": RGBColor(0xC0, 0x39, 0x2B)}


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _kv(doc, pairs):
    for k, v in pairs:
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(v)


def _money_cr(x):
    return f"{x/CR:,.0f}"


def build_docx(report, path: str | Path) -> Path:
    path = Path(path)
    doc = Document()

    # --- title block ------------------------------------------------------
    title = doc.add_heading(f"{report.company_name} ({report.ticker})", level=0)
    sub = doc.add_paragraph("Athena — Equity Research Note")
    sub.runs[0].italic = True
    doc.add_paragraph(f"As of {report.as_of}  |  Currency: {report.currency}")

    # --- recommendation banner -------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Recommendation: {report.recommendation}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = _REC_COLOR.get(report.recommendation, RGBColor(0, 0, 0))
    _kv(doc, [
        ("Market price", f"₹{report.market_price:,.0f}"),
        ("Fair value (blended)", f"₹{report.fair_value:,.0f}"),
        ("Upside / (downside)", f"{report.upside:+.1%}"),
    ])

    # --- investment thesis ------------------------------------------------
    _heading(doc, "Investment Thesis")
    for t in report.thesis:
        doc.add_paragraph(t, style="List Bullet")

    # --- business / financial summary ------------------------------------
    _heading(doc, "Financial Summary")
    _financial_table(doc, report)

    # --- forecast ---------------------------------------------------------
    _heading(doc, "Forecast")
    if "revenue" in report.charts:
        doc.add_picture(report.charts["revenue"], width=Inches(6.0))
    _forecast_table(doc, report)

    # --- DCF --------------------------------------------------------------
    _heading(doc, "DCF Valuation")
    d, w = report.dcf, report.wacc
    _kv(doc, [
        ("Cost of equity (CAPM)", f"{w.cost_of_equity:.2%} "
         f"(Rf {w.risk_free_rate:.2%} + β {w.beta:.2f} × ERP {w.equity_risk_premium:.2%})"),
        ("After-tax cost of debt", f"{w.cost_of_debt_aftertax:.2%}"),
        ("WACC", f"{w.wacc:.2%}  (equity {w.weight_equity:.0%} / debt {w.weight_debt:.0%})"),
        ("PV of explicit FCFF", f"₹{_money_cr(d.pv_explicit)} cr"),
        ("PV of terminal value", f"₹{_money_cr(d.pv_terminal)} cr "
         f"({d.pv_terminal/d.enterprise_value:.0%} of EV)"),
        ("Enterprise value", f"₹{_money_cr(d.enterprise_value)} cr"),
        ("Less: net debt", f"₹{_money_cr(d.net_debt)} cr"),
        ("Equity value", f"₹{_money_cr(d.equity_value)} cr"),
        ("Intrinsic value", f"₹{d.intrinsic_price:,.0f}/share ({d.upside:+.0%} vs market)"),
    ])

    # --- comparables ------------------------------------------------------
    if report.comps is not None:
        _heading(doc, "Comparable Company Analysis")
        _comps_table(doc, report.comps)
        _kv(doc, [("Implied value (blended median multiples)",
                   f"₹{report.comps.implied_price:,.0f}/share "
                   f"({report.comps.upside:+.0%})")])

    # --- Monte Carlo ------------------------------------------------------
    _heading(doc, "Monte Carlo Simulation")
    if "montecarlo" in report.charts:
        doc.add_picture(report.charts["montecarlo"], width=Inches(6.0))
    mc = report.montecarlo
    _kv(doc, [
        ("Mean / median", f"₹{mc.mean:,.0f} / ₹{mc.median:,.0f}"),
        ("Std deviation", f"₹{mc.std:,.0f}"),
        ("90% confidence interval", f"₹{mc.p5:,.0f} – ₹{mc.p95:,.0f}"),
        ("Probability undervalued", f"{mc.prob_undervalued:.0%}"),
    ])

    # --- scenarios --------------------------------------------------------
    _heading(doc, "Bull / Base / Bear Scenarios")
    if "scenarios" in report.charts:
        doc.add_picture(report.charts["scenarios"], width=Inches(4.5))
    _scenario_table(doc, report)

    # --- risks ------------------------------------------------------------
    _heading(doc, "Key Risks")
    _kv(doc, [("Overall risk rating", str(report.risk.overall_level))])
    for r in report.key_risks:
        doc.add_paragraph(r, style="List Bullet")

    # --- recommendation ---------------------------------------------------
    _heading(doc, "Recommendation")
    doc.add_paragraph(
        f"On a blended DCF/peer fair value of ₹{report.fair_value:,.0f} versus a "
        f"market price of ₹{report.market_price:,.0f} ({report.upside:+.1%}), and "
        f"with an overall financial risk rating of {report.risk.overall_level}, "
        f"Athena's model output is {report.recommendation}."
    )
    doc.add_paragraph(
        "This report is generated automatically by Athena for research and "
        "educational purposes only. It is not investment advice.", style="Intense Quote")

    if "sensitivity" in report.charts:
        _heading(doc, "Appendix — Sensitivity")
        doc.add_picture(report.charts["sensitivity"], width=Inches(6.0))

    doc.save(path)
    return path


# --- table helpers ---------------------------------------------------------
def _add_table(doc, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        cell.paragraphs[0].runs[0].bold = True
    return table


def _financial_table(doc, report):
    df = report.fundamentals.sort_index(ascending=False)
    years = list(df.index)
    rows = ["revenue", "ebitda", "ebit", "pat", "total_equity", "debt", "cash"]
    table = _add_table(doc, ["₹ crore"] + [str(y) for y in years])
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = FIELD_LABELS[r]
        for j, y in enumerate(years):
            v = df.loc[y, r]
            cells[j + 1].text = "—" if v != v else f"{v/CR:,.0f}"


def _forecast_table(doc, report):
    fc = report.forecast
    table = _add_table(doc, ["Year", "Revenue", "EBITDA", "EBIT", "FCFF", "Growth"])
    for y, row in fc.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(y)
        cells[1].text = f"{row['revenue']/CR:,.0f}"
        cells[2].text = f"{row['ebitda']/CR:,.0f}"
        cells[3].text = f"{row['ebit']/CR:,.0f}"
        cells[4].text = f"{row['fcff']/CR:,.0f}"
        cells[5].text = f"{row['revenue_growth']*100:.1f}%"


def _comps_table(doc, comps):
    df = comps.table
    table = _add_table(doc, ["Ticker", "EV/Sales", "EV/EBITDA", "P/E", "ROE%", "D/E"])
    for tk, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(tk)
        cells[1].text = f"{row['ev_sales']:.1f}"
        cells[2].text = f"{row['ev_ebitda']:.1f}"
        cells[3].text = f"{row['pe']:.1f}"
        cells[4].text = f"{row['roe']*100:.1f}"
        cells[5].text = f"{row['debt_to_equity']:.2f}"


def _scenario_table(doc, report):
    table = _add_table(doc, ["Scenario", "Intrinsic ₹/share", "Upside",
                             "Growth", "Margin", "WACC", "Term. g"])
    for name, s in report.scenarios.items():
        a = s["assumptions"]
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = f"₹{s['price']:,.0f}"
        cells[2].text = f"{s['upside']:+.0%}"
        cells[3].text = f"{a['growth']*100:.1f}%"
        cells[4].text = f"{a['margin']*100:.1f}%"
        cells[5].text = f"{a['wacc']*100:.1f}%"
        cells[6].text = f"{a['tg']*100:.1f}%"
